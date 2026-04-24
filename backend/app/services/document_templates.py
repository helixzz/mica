from __future__ import annotations

import io
import json
import logging
import re
from decimal import Decimal
from typing import Any
from uuid import UUID

from fastapi import HTTPException
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm import selectinload

from app.core.crypto import decrypt
from app.core.litellm_helpers import resolve_litellm_model
from app.models import (
    AIFeatureRouting,
    AIModel,
    Contract,
    Document,
    DocumentTemplate,
    PaymentSchedule,
    PurchaseOrder,
    Supplier,
)

logger = logging.getLogger(__name__)

PLACEHOLDER_RE = re.compile(r"\[([^\]\n\r]{1,200})\]")

DETERMINISTIC_RESOLVERS: dict[re.Pattern[str], str] = {
    re.compile(r"^PO[_\-\s]?(编号|号|number)$", re.IGNORECASE): "po.po_number",
    re.compile(r"^采购订单[_\-\s]?(编号|号)$"): "po.po_number",
    re.compile(r"^合同[_\-\s]?(编号|号|number)$", re.IGNORECASE): "contract.contract_number",
    re.compile(r"^供应商(名称|名字|name)?$", re.IGNORECASE): "supplier.name",
    re.compile(
        r"^(收款(单位)?|payee)[_\-\s]?(名称|name)?$", re.IGNORECASE
    ): "supplier.payee_name_effective",
    re.compile(r"^((收款(单位)?|payee)[_\-\s]?)?开户行$", re.IGNORECASE): "supplier.payee_bank",
    re.compile(
        r"^(收款(单位)?|银行|bank)[_\-\s]?(账号|账户|account)$", re.IGNORECASE
    ): "supplier.payee_bank_account",
    re.compile(r"^银行账号$"): "supplier.payee_bank_account",
    re.compile(r"^(税号|tax[_\-\s]?number)$", re.IGNORECASE): "supplier.tax_number",
    re.compile(r"^付款期次(说明)?$"): "schedule.label_with_installment",
    re.compile(r"^分期[_\-\s]?(编号|号|no)$", re.IGNORECASE): "schedule.installment_no",
}


def extract_placeholders(content_bytes: bytes | None, filename_template: str) -> list[str]:
    found: list[str] = []
    seen: set[str] = set()

    def _add(text: str) -> None:
        for match in PLACEHOLDER_RE.finditer(text):
            description = match.group(1).strip()
            if description and description not in seen:
                seen.add(description)
                found.append(description)

    _add(filename_template)

    if content_bytes:
        try:
            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content_bytes))
            for para in doc.paragraphs:
                _add(para.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        _add(cell.text)
        except Exception:
            logger.exception("template: failed to parse docx for placeholder extraction")

    return found


def _as_date_str(value: Any) -> str:
    if value is None:
        return ""
    return str(value)


def build_context(
    po: PurchaseOrder,
    contract: Contract,
    supplier: Supplier,
    schedule: PaymentSchedule,
) -> dict[str, Any]:
    payee_name = supplier.payee_name or supplier.name
    amount = (
        schedule.actual_amount if schedule.actual_amount is not None else schedule.planned_amount
    )
    date_value = schedule.actual_date or schedule.planned_date
    return {
        "po": {
            "po_number": po.po_number,
            "currency": po.currency,
            "total_amount": str(po.total_amount),
            "status": po.status,
        },
        "contract": {
            "contract_number": contract.contract_number,
            "title": contract.title,
            "total_amount": str(contract.total_amount),
            "currency": contract.currency,
            "signed_date": _as_date_str(contract.signed_date),
            "effective_date": _as_date_str(contract.effective_date),
            "expiry_date": _as_date_str(contract.expiry_date),
            "status": contract.status,
        },
        "supplier": {
            "name": supplier.name,
            "code": supplier.code,
            "tax_number": supplier.tax_number or "",
            "contact_name": supplier.contact_name or "",
            "contact_phone": supplier.contact_phone or "",
            "contact_email": supplier.contact_email or "",
            "payee_name": supplier.payee_name or "",
            "payee_name_effective": payee_name,
            "payee_bank": supplier.payee_bank or "",
            "payee_bank_account": supplier.payee_bank_account or "",
        },
        "schedule": {
            "installment_no": schedule.installment_no,
            "label": schedule.label,
            "label_with_installment": f"第 {schedule.installment_no} 期 · {schedule.label}",
            "planned_amount": str(schedule.planned_amount),
            "planned_date": _as_date_str(schedule.planned_date),
            "actual_amount": str(schedule.actual_amount) if schedule.actual_amount else "",
            "actual_date": _as_date_str(schedule.actual_date),
            "status": schedule.status,
            "effective_amount": str(amount),
            "effective_date": _as_date_str(date_value),
            "trigger_type": schedule.trigger_type,
            "trigger_description": schedule.trigger_description or "",
        },
    }


def _lookup_context(context: dict[str, Any], path: str) -> str | None:
    cur: Any = context
    for key in path.split("."):
        if not isinstance(cur, dict) or key not in cur:
            return None
        cur = cur[key]
    return None if cur is None else str(cur)


def resolve_placeholder_deterministic(description: str, context: dict[str, Any]) -> str | None:
    for pattern, path in DETERMINISTIC_RESOLVERS.items():
        if pattern.search(description):
            value = _lookup_context(context, path)
            if value is not None:
                return value
    return None


async def resolve_placeholders_with_llm(
    db: AsyncSession,
    placeholders: list[str],
    context: dict[str, Any],
) -> dict[str, str]:
    if not placeholders:
        return {}

    routing = (
        await db.execute(
            select(AIFeatureRouting).where(AIFeatureRouting.feature_code == "document_generation")
        )
    ).scalar_one_or_none()
    model: AIModel | None = None
    if routing and routing.enabled and routing.primary_model_id:
        model = (
            await db.execute(select(AIModel).where(AIModel.id == routing.primary_model_id))
        ).scalar_one_or_none()

    if model is None:
        return {}

    try:
        from litellm import acompletion
    except ImportError:
        return {}

    prompt = _build_llm_prompt(placeholders, context)
    kwargs: dict[str, Any] = {
        "model": resolve_litellm_model(model.provider, model.model_string),
        "messages": [
            {
                "role": "system",
                "content": "You are a precise data-extraction assistant. "
                "Given structured business data and a list of placeholder descriptions, "
                "return ONLY a valid JSON object mapping each exact placeholder text to "
                "its extracted value. No prose, no markdown fences.",
            },
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.0,
        "max_tokens": 1200,
    }
    api_key = decrypt(model.api_key_encrypted) if model.api_key_encrypted else None
    if api_key:
        kwargs["api_key"] = api_key
    if model.api_base:
        kwargs["api_base"] = model.api_base

    try:
        resp = await acompletion(**kwargs)
        raw = resp.choices[0].message.content or "{}"
        raw = raw.strip()
        if raw.startswith("```"):
            raw = raw.strip("`").lstrip("json").strip()
        data = json.loads(raw)
        if not isinstance(data, dict):
            return {}
        return {str(k): str(v) if v is not None else "" for k, v in data.items()}
    except Exception:
        logger.exception("template: LLM placeholder resolution failed")
        return {}


def _build_llm_prompt(placeholders: list[str], context: dict[str, Any]) -> str:
    return (
        "Here is structured business data about a purchase order, a contract, "
        "a supplier, and a single payment schedule installment:\n\n"
        f"{json.dumps(context, ensure_ascii=False, indent=2, default=str)}\n\n"
        "For each of the following placeholders, produce the best-fitting extracted "
        "or formatted string value. If a placeholder description implies a date format "
        '(e.g. "YYYY-MM-DD" or "YYYY年MM月DD日"), format accordingly. '
        'If it implies an amount in words (大写), convert the number to Chinese "大写金额" '
        '(e.g. "肆佰伍拾万元整"). If you cannot find a value, return an empty string.\n\n'
        "Return ONLY a single JSON object, where keys are the EXACT placeholder texts "
        "as given, and values are the resolved strings.\n\n"
        "Placeholders:\n" + "\n".join(f"- {p}" for p in placeholders)
    )


_DIGIT_CN = "零壹贰叁肆伍陆柒捌玖"
_UNIT_CN = ["", "拾", "佰", "仟"]
_BIG_UNIT_CN = ["", "万", "亿", "兆"]


def cn_amount_upper(amount: Decimal | float | int | str | None) -> str:
    if amount is None or amount == "":
        return ""
    try:
        d = Decimal(str(amount)).quantize(Decimal("0.01"))
    except Exception:
        return str(amount)
    yuan_part = int(d)
    cents = int((d - Decimal(yuan_part)) * 100)

    def _convert_int(n: int) -> str:
        if n == 0:
            return "零"
        parts: list[str] = []
        big = 0
        while n > 0:
            section = n % 10000
            n //= 10000
            if section == 0:
                parts.insert(0, "")
                big += 1
                continue
            section_str = ""
            prev_zero = False
            for i in range(4):
                digit = section % 10
                section //= 10
                if digit == 0:
                    if not prev_zero and section_str:
                        section_str = "零" + section_str
                    prev_zero = True
                else:
                    section_str = _DIGIT_CN[digit] + _UNIT_CN[i] + section_str
                    prev_zero = False
            section_str = section_str.rstrip("零")
            parts.insert(0, section_str + _BIG_UNIT_CN[big])
            big += 1
        return "".join(p for p in parts if p).lstrip("零") or "零"

    yuan_str = _convert_int(yuan_part) + "元"
    if cents == 0:
        return yuan_str + "整"
    jiao = cents // 10
    fen = cents % 10
    tail = ""
    if jiao:
        tail += _DIGIT_CN[jiao] + "角"
    if fen:
        tail += _DIGIT_CN[fen] + "分"
    else:
        tail += "整" if jiao else "零"
    return yuan_str + tail


DATE_FORMAT_RE = re.compile(
    r"(YYYY|YY)[-/\.\s年]*MM[-/\.\s月]*DD[日]?|(YYYY|YY)MM?DD?",
    re.IGNORECASE,
)


def _format_schedule_date(fmt: str, y: str, m: str, d: str) -> str:
    out = fmt
    out = out.replace("YYYY", y).replace("yyyy", y)
    out = out.replace("YY", y[-2:]).replace("yy", y[-2:])
    out = out.replace("MM", m).replace("mm", m)
    out = out.replace("DD", d).replace("dd", d)
    return out


def _enrich_with_computed(
    description: str, context: dict[str, Any], current: str | None
) -> str | None:
    if current:
        return current
    if "大写" in description or "upper" in description.lower():
        amount = context["schedule"]["effective_amount"]
        return cn_amount_upper(amount)
    if "数字" in description and ("金额" in description or "amount" in description.lower()):
        return context["schedule"]["effective_amount"]
    date_match = DATE_FORMAT_RE.search(description)
    if date_match:
        raw_date = context["schedule"]["effective_date"]
        if not raw_date or raw_date == "None":
            return ""
        try:
            y, m, d = raw_date.split("-")
        except ValueError:
            return raw_date
        return _format_schedule_date(date_match.group(0), y, m, d)
    return None


async def resolve_all_placeholders(
    db: AsyncSession,
    placeholders: list[str],
    context: dict[str, Any],
) -> dict[str, str]:
    resolved: dict[str, str] = {}
    deterministic: list[str] = []
    unresolved: list[str] = []

    for p in placeholders:
        value = resolve_placeholder_deterministic(p, context)
        if value is not None:
            resolved[p] = value
            deterministic.append(p)
        else:
            unresolved.append(p)

    if unresolved:
        llm_results = await resolve_placeholders_with_llm(db, unresolved, context)
        for p in unresolved:
            resolved[p] = llm_results.get(p, "")

    for p in placeholders:
        enriched = _enrich_with_computed(p, context, resolved.get(p) or None)
        if enriched is not None:
            resolved[p] = enriched

    return resolved


def render_filename(filename_template: str, mapping: dict[str, str]) -> str:
    def _sub(match: re.Match[str]) -> str:
        return mapping.get(match.group(1).strip(), "")

    filename = PLACEHOLDER_RE.sub(_sub, filename_template)
    sanitized = re.sub(r"[<>:\"/\\|?*\x00-\x1f]+", "_", filename)
    return sanitized.strip()


def substitute_docx(content_bytes: bytes, mapping: dict[str, str]) -> bytes:
    from docx import Document as DocxDocument

    doc = DocxDocument(io.BytesIO(content_bytes))

    def _replace_in_paragraph(para) -> None:
        full_text = "".join(run.text for run in para.runs)
        if not PLACEHOLDER_RE.search(full_text):
            return
        new_text = PLACEHOLDER_RE.sub(lambda m: mapping.get(m.group(1).strip(), ""), full_text)
        if para.runs:
            para.runs[0].text = new_text
            for run in para.runs[1:]:
                run.text = ""
        else:
            para.text = new_text

    for para in doc.paragraphs:
        _replace_in_paragraph(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generate_payment_document(
    db: AsyncSession,
    template_code: str,
    schedule_item_id: UUID,
) -> tuple[bytes, str]:
    template = (
        await db.execute(
            select(DocumentTemplate)
            .where(DocumentTemplate.code == template_code)
            .options(selectinload(DocumentTemplate.template_document))
        )
    ).scalar_one_or_none()
    if template is None:
        raise HTTPException(404, "template.not_found")
    if not template.is_enabled:
        raise HTTPException(409, "template.disabled")
    if template.template_document is None:
        raise HTTPException(409, "template.no_file")

    schedule = (
        await db.execute(
            select(PaymentSchedule)
            .where(PaymentSchedule.id == schedule_item_id)
            .options(
                selectinload(PaymentSchedule.contract).selectinload(Contract.po),
                selectinload(PaymentSchedule.contract).selectinload(Contract.supplier),
                selectinload(PaymentSchedule.po).selectinload(PurchaseOrder.supplier),
            )
        )
    ).scalar_one_or_none()
    if schedule is None:
        raise HTTPException(404, "schedule_item.not_found")

    if schedule.contract is not None:
        contract = schedule.contract
        po = contract.po
        supplier = contract.supplier
    elif schedule.po is not None:
        po = schedule.po
        contract = (
            await db.execute(select(Contract).where(Contract.po_id == po.id).limit(1))
        ).scalar_one_or_none()
        if contract is None:
            raise HTTPException(
                409,
                "template.contract_required_for_generation",
            )
        supplier = (
            await db.execute(select(Supplier).where(Supplier.id == po.supplier_id))
        ).scalar_one_or_none()
    else:
        raise HTTPException(409, "schedule_item.orphan")

    if supplier is None:
        raise HTTPException(409, "template.supplier_missing")

    file_bytes = _read_document_bytes(template.template_document)

    context = build_context(po, contract, supplier, schedule)
    placeholders = extract_placeholders(file_bytes, template.filename_template)
    mapping = await resolve_all_placeholders(db, placeholders, context)

    generated_bytes = substitute_docx(file_bytes, mapping)
    filename_base = render_filename(template.filename_template, mapping)
    if not filename_base:
        filename_base = f"{template.code}_{schedule.installment_no}"
    filename = f"{filename_base}.docx"

    return generated_bytes, filename


def _read_document_bytes(document: Document) -> bytes:
    from pathlib import Path

    from app.config import get_settings

    settings = get_settings()
    media_root = Path(getattr(settings, "media_root", "/app/media"))
    path = media_root / document.storage_key
    if not path.exists():
        raise HTTPException(500, "template.file_missing_on_disk")
    return path.read_bytes()
